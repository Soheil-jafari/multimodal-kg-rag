"""Export the experiment log to .docx as the supporting appendix.

    python -m scripts.export_log_docx

Structure is preserved rather than flattened: headings stay headings, tables stay
tables, and fenced blocks plus inline code keep a monospace face — a log whose tables
have collapsed into prose is no longer evidence of anything.

Covers exactly the constructs the log uses (surveyed, not guessed): H1-H3, horizontal
rules, blockquotes, bullet and numbered lists, pipe tables, fenced blocks, and inline
bold/italic/code. An unhandled construct would silently lose its formatting, so anything
unrecognised falls through as a plain paragraph with its markers intact rather than
disappearing.
"""
from __future__ import annotations

import argparse
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
BODY, MONO = "Calibri", "Consolas"
INK = RGBColor(0x14, 0x18, 0x1D)
MUTED = RGBColor(0x3D, 0x46, 0x50)

#: **bold** / *italic* / `code`, longest-first so ** is consumed before *
_TOK = re.compile(r"(\*\*[^*]+\*\*|(?<!\*)\*(?!\*)[^*]+\*(?!\*)|`[^`]+`)")


def add_code(par, block: list, size=8.6) -> None:
    """A fenced block: monospace, real line breaks, and NO inline parsing.

    Markdown emphasis characters inside a config excerpt or captured output are literal,
    so running them through the inline tokenizer would eat asterisks and backticks that
    are part of the recorded text. Line breaks are real break elements — a newline inside
    a run is not valid document XML.
    """
    for k, ln in enumerate(block):
        r = par.add_run(ln)
        r.font.name = MONO
        r.font.size = Pt(size)
        r.font.color.rgb = INK
        if k < len(block) - 1:
            r.add_break()


def add_runs(par, text: str, size=9.5, mono_all=False) -> None:
    """Emit inline runs, honouring bold / italic / code."""
    for part in _TOK.split(text):
        if not part:
            continue
        bold = italic = code = False
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            part, bold = part[2:-2], True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            part, code = part[1:-1], True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            part, italic = part[1:-1], True
        r = par.add_run(part)
        r.bold, r.italic = bold, italic
        r.font.name = MONO if (code or mono_all) else BODY
        r.font.size = Pt(size - 0.7 if code or mono_all else size)
        r.font.color.rgb = MUTED if code else INK


def shade(cell, hex_fill: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def hrule(doc) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    bd = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "1"), ("w:color", "B9C2CC")):
        bottom.set(qn(k), v)
    bd.append(bottom)
    pPr.append(bd)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)


def cells_of(row: str) -> list:
    return [c.strip() for c in row.strip().strip("|").split("|")]


def add_table(doc, rows: list) -> None:
    head = cells_of(rows[0])
    body = [cells_of(r) for r in rows[2:]]          # row 1 is the alignment separator
    width = max([len(head)] + [len(r) for r in body])
    t = doc.add_table(rows=1, cols=width)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(head + [""] * (width - len(head))):
        cell = t.rows[0].cells[i]
        cell.text = ""
        add_runs(cell.paragraphs[0], c, size=8.5)
        for r in cell.paragraphs[0].runs:
            r.bold = True
        shade(cell, "E7ECF1")
    for br in body:
        cells = t.add_row().cells
        for i, c in enumerate(br + [""] * (width - len(br))):
            cells[i].text = ""
            add_runs(cells[i].paragraphs[0], c, size=8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def convert(md_path: str, out_path: str) -> dict:
    lines = open(md_path, encoding="utf-8").read().split("\n")
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name, st.font.size = BODY, Pt(9.5)
    st.paragraph_format.space_after = Pt(4)
    for s in doc.sections:
        s.left_margin = s.right_margin = Pt(54)
        s.top_margin = s.bottom_margin = Pt(48)

    seen = {k: 0 for k in ("h1", "h2", "h3", "table", "code", "quote", "list", "para", "hr")}
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), 0); seen["h1"] += 1; i += 1; continue
        if line.startswith("## "):
            doc.add_page_break() if seen["h2"] else None
            doc.add_heading(line[3:].strip(), 1); seen["h2"] += 1; i += 1; continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), 2); seen["h3"] += 1; i += 1; continue

        if line.startswith("```"):
            block, i = [], i + 1
            while i < n and not lines[i].startswith("```"):
                block.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            p.paragraph_format.space_before = Pt(3)
            add_code(p, block)
            seen["code"] += 1
            continue

        if line.startswith("|"):
            rows = []
            while i < n and lines[i].startswith("|"):
                rows.append(lines[i]); i += 1
            if len(rows) >= 2:
                add_table(doc, rows); seen["table"] += 1
            continue

        if line.startswith("> "):
            block = []
            while i < n and lines[i].startswith(">"):
                block.append(lines[i].lstrip(">").strip()); i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.space_before = Pt(4)
            pPr = p._p.get_or_add_pPr()
            bd = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            for k, v in (("w:val", "single"), ("w:sz", "18"), ("w:space", "8"),
                         ("w:color", "8B96A2")):
                left.set(qn(k), v)
            bd.append(left)
            pPr.append(bd)
            add_runs(p, " ".join(x for x in block if x))
            seen["quote"] += 1
            continue

        if re.match(r"^(\s*)([-*] |\d+\. )", line):
            while i < n and re.match(r"^(\s*)([-*] |\d+\. )", lines[i]):
                raw = lines[i]
                nested = len(raw) - len(raw.lstrip())
                item = re.sub(r"^\s*([-*] |\d+\. )", "", raw)
                numbered = bool(re.match(r"^\s*\d+\. ", raw))
                i += 1
                while i < n and lines[i].startswith("  ") and lines[i].strip() \
                        and not re.match(r"^(\s*)([-*] |\d+\. )", lines[i]):
                    item += " " + lines[i].strip(); i += 1
                p = doc.add_paragraph(
                    style="List Number" if numbered else "List Bullet")
                if nested:
                    p.paragraph_format.left_indent = Pt(36)
                add_runs(p, item.strip())
            seen["list"] += 1
            continue

        if line.strip() == "---":
            hrule(doc); seen["hr"] += 1; i += 1; continue
        if not line.strip():
            i += 1; continue

        para = []
        while i < n and lines[i].strip() and not re.match(
                r"^(#{1,3} |\||> |```|\s*[-*] |\s*\d+\. )", lines[i]) and lines[i].strip() != "---":
            para.append(lines[i].strip()); i += 1
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, " ".join(para))
        seen["para"] += 1

    doc.save(out_path)
    return seen


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--md", default=ROOT + "/artifacts/EXPERIMENT_LOG.md")
    ap.add_argument("--out", default=ROOT + "/reports/EXPERIMENT_LOG.docx")
    args = ap.parse_args()

    seen = convert(args.md, args.out)
    print(f"wrote {os.path.relpath(args.out, ROOT)}  "
          f"({os.path.getsize(args.out)/1024:.0f} KB)")
    print("  converted: " + "  ".join(f"{k}={v}" for k, v in seen.items()))


if __name__ == "__main__":
    main()
